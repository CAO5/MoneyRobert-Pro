from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "final"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
RED = "9B1C1C"
GOLD = "7A5A00"
BLACK = "1F1F1F"
FONT_LATIN = "Calibri"
FONT_CN = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        s = styles[name]
        s.font.name = FONT_LATIN
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = FONT_LATIN
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.5)
        s.paragraph_format.first_line_indent = Inches(-0.25)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(running_title), 9, False, MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(p.add_run("MoneyRobert Pro  |  "), 9, False, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def cover(doc, title, subtitle, doc_code, version="V2.0"):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(title), 28, True, DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    set_run_font(p.add_run(subtitle), 14, False, MUTED)
    meta = [
        ("文档编号", doc_code),
        ("版本", version),
        ("文档状态", "正式设计基线"),
        ("适用系统", "MoneyRobert Pro"),
        ("编制日期", str(date.today())),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for i, (k, v) in enumerate(meta):
        set_cell_shading(table.cell(i, 0), LIGHT_GRAY)
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for c in table.rows[i].cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, 10.5, c is table.cell(i, 0))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("内部设计与开发使用"), 10, True, GOLD)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), 11, True)
        set_run_font(p.add_run(text[len(bold_prefix):]), 11)
    else:
        set_run_font(p.add_run(text), 11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item), 11)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(item), 11)


def add_note(doc, title, text, kind="info"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    fill = LIGHT_BLUE if kind == "info" else ("FFF4CE" if kind == "warn" else "FDE7E9")
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{title}："), 10.5, True, GOLD if kind == "warn" else (RED if kind == "risk" else DARK_BLUE))
    set_run_font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 and len(headers) <= 4 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, font_size, ri == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_doc_control(doc, purpose, audience):
    add_heading(doc, "文档说明", 1)
    add_table(doc, ["项目", "内容"], [
        ("编制目的", purpose),
        ("目标读者", audience),
        ("状态口径", "【现状】已在代码中实现；【修复】已实现但必须整改；【规划】目标版本新增能力。"),
        ("优先级", "P0 阻断实盘；P1 核心能力；P2 增强能力；P3 体验优化。"),
    ], [2000, 7360], 9.5)
    add_heading(doc, "目录", 1)


def add_contents(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(item), 11)
    doc.add_page_break()


def requirement_document():
    doc = Document()
    style_document(doc, "MoneyRobert Pro · 软件需求规格说明书")
    cover(doc, "软件需求规格说明书", "面向加密资产分析、模拟交易与风险控制平台", "MRP-SRS-002")
    add_doc_control(doc, "定义系统目标、范围、角色、功能、数据、模型、安全、非功能需求和验收标准，作为产品、研发、测试和上线评审共同基线。", "产品经理、经济/量化研究员、基金经理、架构师、研发、测试、运维与合规人员")
    add_contents(doc, ["项目概述", "业务目标与成功指标", "系统边界与版本范围", "用户角色与核心场景", "功能需求", "量化与经济模型需求", "资金与风险管理需求", "数据需求", "非功能需求", "安全与合规需求", "验收标准", "实施优先级与版本路线"])

    add_heading(doc, "1 项目概述", 1)
    add_p(doc, "MoneyRobert Pro 是一个面向加密资产市场的研究、决策、模拟、回测与交易辅助平台。系统通过市场数据、技术指标、衍生品资金数据、新闻事件和多 Agent 分析生成候选交易观点，并由确定性风控与执行模块约束交易行为。")
    add_note(doc, "产品定位", "系统输出的是带概率、成本和风险约束的决策支持，不承诺稳定盈利，也不得以 LLM 自信程度替代统计概率。")
    add_heading(doc, "1.1 建设原则", 2)
    add_bullets(doc, ["先保证账本、数据与回测可信，再扩大自动化权限。", "预测、组合、风控和执行分层；LLM 只生成候选观点和解释。", "每项结论必须可追溯到数据时间、模型版本、规则版本和历史表现。", "模拟盘、回测和实盘共用一致的订单、费用、仓位和风险语义。", "默认保护用户资金，风险不确定时输出观望。"])

    add_heading(doc, "2 业务目标与成功指标", 1)
    add_table(doc, ["目标", "度量指标", "目标值/门槛", "优先级"], [
        ("研究可信", "预测校准、数据新鲜度、样本外表现", "概率校准误差可量化；关键数据有新鲜度状态", "P0/P1"),
        ("交易可信", "账户资金守恒、订单幂等、费用完整", "自动化测试100%覆盖关键资金路径", "P0"),
        ("风险可控", "最大回撤、CVaR、风险预算超限", "超限阻断且全量审计", "P0"),
        ("决策透明", "证据覆盖率、反对证据、模型卡", "每个建议包含数据来源与失效条件", "P1"),
        ("操作高效", "从发现机会到模拟执行的步骤数", "核心任务5步内完成", "P2"),
    ], [1700, 2800, 3400, 1460], 9)

    add_heading(doc, "3 系统边界与版本范围", 1)
    add_table(doc, ["状态", "范围"], [
        ("【现状】", "OKX 行情与交易、K线/资金费率、技术分析、新闻、Agent 辩论、模拟盘、自动交易配置、回测、通知、报告、记忆与演进。"),
        ("【修复】", "回测撮合和现金记账、绩效指标、测试接入、报告数据隔离、数据质量告警、错误吞噬。"),
        ("【规划】", "特征仓库、概率预测、订单簿/逐笔/清算/期权/链上/宏观数据、组合风险、模型治理、可靠任务队列。"),
        ("不在范围", "保证收益、无人工治理的无限制实盘、绕过交易所或地区监管、将生成式文本直接作为订单。"),
    ], [1800, 7560])

    add_heading(doc, "4 用户角色与核心场景", 1)
    add_table(doc, ["角色", "主要任务", "关键权限"], [
        ("普通投资者", "查看市场、理解建议、模拟交易、管理风险预算", "仅访问本人账户与报告"),
        ("专业交易员", "创建策略、回测、纸面/模拟/实盘执行", "受风险等级与审批约束"),
        ("基金经理", "组合配置、资金分配、绩效归因、策略停用", "组合级风险与执行审批"),
        ("量化研究员", "特征、模型、验证、校准、模型版本管理", "不可直接绕过执行风控"),
        ("风控人员", "设置限额、查看暴露、熔断与恢复", "独立否决与紧急停止"),
        ("系统管理员", "用户、密钥、数据源、模型供应商、审计", "高危操作二次确认"),
    ], [1700, 4400, 3260], 9.2)

    add_heading(doc, "5 功能需求", 1)
    functional_groups = [
        ("FR-MD 市场数据", [
            ("FR-MD-001", "采集 ticker、K线、资金费率并记录源时间、接收时间和质量状态。", "P0"),
            ("FR-MD-002", "支持持仓量、多空比、订单簿、逐笔成交、清算、基差和跨交易所数据扩展。", "P1"),
            ("FR-MD-003", "检测缺口、重复、异常价格、时钟漂移和数据过期；过期数据不得触发实盘。", "P0"),
            ("FR-MD-004", "交易对、周期、保留期和采集频率必须配置化。", "P1"),
        ]),
        ("FR-AN 分析与预测", [
            ("FR-AN-001", "展示趋势、动量、波动、量价、资金和情绪特征，不把单一指标作为结论。", "P1"),
            ("FR-AN-002", "预测输出上涨/下跌/震荡概率、收益分位数、预期波动率和不确定度。", "P1"),
            ("FR-AN-003", "每次预测记录特征、模型、训练窗口、预测周期和校准版本。", "P1"),
            ("FR-AN-004", "缺少关键数据或模型失配时必须输出观望与原因。", "P0"),
        ]),
        ("FR-AG 多 Agent 决策", [
            ("FR-AG-001", "技术、资金、新闻、宏观、链上等 Agent 形成独立意见，并标记数据来源。", "P1"),
            ("FR-AG-002", "硬编码兜底意见不得作为有效共识，仅可标记为不可用。", "P0"),
            ("FR-AG-003", "基金经理聚合时使用历史可靠度和概率校准，不直接使用文本自信度。", "P1"),
            ("FR-AG-004", "保留支持、反对、冲突、最终否决和决策证据。", "P1"),
        ]),
        ("FR-BT 回测与模拟", [
            ("FR-BT-001", "支持市价、限价、部分成交、加减仓、反手、止盈止损、费用、滑点和资金费率。", "P0"),
            ("FR-BT-002", "任何成交后满足现金+持仓市值=权益的会计恒等式。", "P0"),
            ("FR-BT-003", "采用 walk-forward 和 purged validation，阻止未来数据泄漏。", "P1"),
            ("FR-BT-004", "输出净收益、回撤、Sharpe、Sortino、Calmar、CVaR、换手率和基准比较。", "P1"),
        ]),
        ("FR-EX 交易执行", [
            ("FR-EX-001", "支持纸面、交易所模拟和实盘三级环境，界面持续显示当前环境。", "P0"),
            ("FR-EX-002", "订单具备客户端幂等键、状态机、重试边界和交易所回查。", "P0"),
            ("FR-EX-003", "实盘执行前必须通过数据、模型、组合和账户四层风控。", "P0"),
            ("FR-EX-004", "紧急停止立即阻止新订单，并给出持仓处置选项。", "P0"),
        ]),
    ]
    for title, rows in functional_groups:
        add_heading(doc, title, 2)
        add_table(doc, ["编号", "需求说明", "优先级"], rows, [1600, 6500, 1260], 9.2)

    add_heading(doc, "6 量化与经济模型需求", 1)
    add_p(doc, "系统应采用异质预期与市场状态相结合的框架：不同信息主体形成条件观点，模型根据市场状态、历史可靠度和数据质量聚合观点。经济模型不是装饰性 Agent 名称，而必须映射到可观测变量。")
    add_table(doc, ["模型域", "输入", "输出/用途"], [
        ("宏观流动性", "实际利率、美元、流动性、稳定币、ETF流量", "风险偏好与风险溢价状态"),
        ("衍生品期限结构", "资金费率、基差、持仓量、期权IV/Skew", "拥挤度、carry与尾部风险"),
        ("市场微观结构", "spread、深度、订单流、冲击成本", "可执行价格和容量"),
        ("状态切换", "趋势、波动、相关性、流动性", "趋势/震荡/危机等 regime"),
        ("概率预测", "结构化特征与状态", "概率、分位数和不确定度"),
    ], [2100, 3900, 3360], 9.2)
    add_note(doc, "数学期望门槛", "仅当扣除手续费、滑点、资金费率和冲击成本后的条件期望为正，且 CVaR 位于风险预算内，系统才可生成可执行候选信号。")

    add_heading(doc, "7 资金与风险管理需求", 1)
    add_bullets(doc, ["仓位以单笔风险预算、止损距离和波动率计算，不以 Agent 置信度线性放大杠杆。", "组合层限制单品种、单策略、单方向、相关资产簇和总杠杆暴露。", "支持 Fractional Kelly，但 Kelly 输入必须是样本外校准概率，并设置硬上限。", "监控日/周/月回撤、CVaR、连续亏损、流动性和交易频率。", "风控参数变更需要版本、操作者、原因和审批记录。"])

    add_heading(doc, "8 数据需求", 1)
    add_table(doc, ["数据层", "要求", "保留策略"], [
        ("原始层", "保存交易所原始响应或可重放事件，禁止覆盖", "研究数据长期保留"),
        ("标准层", "统一 symbol、时间、价格精度、币种和来源", "按数据等级分区"),
        ("特征层", "point-in-time 正确，包含计算版本", "支持回溯重算"),
        ("预测层", "保存输入快照、输出、模型和校准信息", "与审计周期一致"),
        ("交易账本", "订单、成交、费用、资金费率、仓位、权益不可抵赖", "永久或按法规要求"),
    ], [1700, 4800, 2860], 9.2)

    add_heading(doc, "9 非功能需求", 1)
    add_table(doc, ["类别", "要求"], [
        ("可靠性", "关键订单路径可恢复；服务重启不得丢失订单状态；数据库操作幂等。"),
        ("性能", "实时行情P95端到端延迟可观测；核心查询P95小于500ms（不含外部LLM）。"),
        ("扩展性", "采集、特征、模型和执行通过稳定接口解耦，支持新增交易所和模型。"),
        ("可观测性", "指标、日志、追踪和业务审计关联同一 correlation_id。"),
        ("可维护性", "路由层不承载模型算法；核心模块具有单元、属性和集成测试。"),
        ("可用性", "关键页面明确加载、空数据、过期、降级和错误状态。"),
    ], [1900, 7460], 9.5)

    add_heading(doc, "10 安全与合规需求", 1)
    add_bullets(doc, ["所有用户资源必须按 user_id/tenant_id 隔离，敏感表建议启用 PostgreSQL RLS。", "API Key 仅加密存储，密钥轮换，日志严禁输出明文。", "JWT 区分 access/refresh 类型，支持吊销、设备会话和最小权限。", "实盘、风控放宽、密钥变更、紧急停止恢复均需二次确认和审计。", "新闻和外部文本均视为不可信输入，防止 Prompt Injection。", "产品页面持续声明模型局限、样本量和非投资建议边界。"])

    add_heading(doc, "11 验收标准", 1)
    add_table(doc, ["验收域", "必须通过的标准"], [
        ("财务正确性", "随机订单序列下资金守恒；反手、部分成交、费用和资金费率测试通过。"),
        ("回测可信", "无未来函数；使用日末权益；所有成本计入；可复现实验结果。"),
        ("数据可信", "缺口、过期和异常可见；过期数据阻断实盘。"),
        ("模型可信", "有样本外基准、校准曲线、模型卡和失效条件。"),
        ("风险可信", "任何订单均不能绕过组合风控；紧急停止演练通过。"),
        ("安全可信", "越权测试、密钥保护、审计完整性和依赖扫描通过。"),
    ], [2200, 7160], 9.5)

    add_heading(doc, "12 实施优先级与版本路线", 1)
    add_table(doc, ["阶段", "目标", "主要交付"], [
        ("R1 可信底座", "阻断错误资金结果", "撮合修复、资金守恒测试、绩效修复、数据隔离、全测试接入"),
        ("R2 数据与研究", "形成可复现实验", "长期数据、特征仓库、数据质量、walk-forward、基准模型"),
        ("R3 组合风控", "从单信号升级到组合", "风险预算、相关性、CVaR、流动性、统一账本"),
        ("R4 可信体验", "让用户理解而非盲信", "概率决策卡、证据血缘、模型卡、交易归因"),
        ("R5 受控自动化", "有限权限实盘", "影子模式、审批晋级、可靠执行、持续监控"),
    ], [1700, 2600, 5060], 9.2)
    path = OUT / "MoneyRobert-Pro_软件需求规格说明书_V2.0.docx"
    doc.save(path)
    return path


def design_document():
    doc = Document()
    style_document(doc, "MoneyRobert Pro · 开发设计文档")
    cover(doc, "开发设计文档", "目标架构、领域模型、量化模型、数据、接口、安全与测试设计", "MRP-SDD-002")
    add_doc_control(doc, "给出系统从当前模块化单体演进为可信交易分析平台的详细设计，明确组件职责、关键算法、数据结构、接口、部署、测试和迁移方案。", "架构师、后端/前端/数据/量化研发、测试、运维、安全与技术负责人")
    add_contents(doc, ["设计目标与约束", "现状架构评估", "目标逻辑架构", "领域与模块设计", "数据架构", "量化预测设计", "组合与风险设计", "回测与执行设计", "Agent与LLM设计", "接口与事件设计", "前端设计", "安全设计", "可观测性与运维", "测试策略", "迁移与发布方案"])

    add_heading(doc, "1 设计目标与约束", 1)
    add_p(doc, "系统采用模块化单体作为近期部署形态，以领域边界和事件契约实现逻辑解耦。只有在吞吐、团队边界或独立扩缩容产生明确需求时再拆分服务。")
    add_bullets(doc, ["核心账本与风控必须确定性、可测试、可重放。", "所有研究输入满足 point-in-time 语义。", "预测输出为概率分布，不输出未经校准的确定性价格结论。", "实盘执行与 LLM 隔离，任何模型不可直接调用交易所下单。", "以 PostgreSQL 为事实源，Redis 只承担缓存、锁和短期队列。"])

    add_heading(doc, "2 现状架构评估", 1)
    add_table(doc, ["方面", "现状", "设计决策"], [
        ("后端", "Rust/Axum 模块化单体，路由模块较多", "保留单体，拆出 application/domain/infrastructure 层"),
        ("分析", "ai_analysis.rs 聚合HTTP、指标、Prompt和评分", "拆为分析用例、特征服务、Agent编排、评分策略"),
        ("行情", "20交易对串行REST轮询", "WebSocket优先、REST补偿、并发限速与质量状态"),
        ("回测", "已有撮合、账户、风险、绩效模块", "先修账本，再统一回测/模拟/实盘语义"),
        ("前端", "Vue页面功能完整但大型页面集中", "按领域拆页面容器、组件、查询与状态"),
        ("任务", "进程内异步任务", "持久化Job、租约、重试、幂等与恢复"),
    ], [1700, 3600, 4060], 9.2)
    add_note(doc, "P0 技术债", "撮合引擎的平仓现金处理和超量反手逻辑必须在任何模型评估、Agent晋级或实盘扩权之前修复。", "risk")

    add_heading(doc, "3 目标逻辑架构", 1)
    add_p(doc, "推荐数据流：Market Data → Normalization → Feature Store → Model/Signal → Portfolio/Risk → Execution → Ledger/Attribution。控制流与数据流分离，所有阶段使用 correlation_id 和版本信息连接。")
    add_table(doc, ["层", "核心组件", "职责"], [
        ("接入层", "REST/WebSocket API、认证、限流", "输入校验、身份、协议转换，不实现业务算法"),
        ("应用层", "Use Case、Job Orchestrator", "事务边界、流程编排、权限与幂等"),
        ("领域层", "Market、Signal、Portfolio、Risk、Execution、Ledger", "纯业务规则与状态机"),
        ("研究层", "Feature、Model、Calibration、Backtest", "特征计算、预测、验证和模型治理"),
        ("基础设施", "PostgreSQL、Redis、OKX、LLM、消息/任务", "外部适配、持久化、缓存和可靠交付"),
        ("体验层", "Vue Query/Store、决策卡、交易台", "任务流程、证据展示和风险交互"),
    ], [1500, 3000, 4860], 9.2)

    add_heading(doc, "4 领域与模块设计", 1)
    domains = [
        ("Market Data", "Instrument、Tick、Bar、OrderBook、Trade、Funding、OpenInterest", "数据标准化和质量状态"),
        ("Research", "FeatureSet、DatasetSnapshot、Experiment、ModelVersion", "可重复研究与防泄漏"),
        ("Decision", "Forecast、AgentOpinion、DecisionEvidence、TradeCandidate", "概率观点和证据聚合"),
        ("Portfolio", "Portfolio、Position、Exposure、RiskBudget", "组合目标与风险分配"),
        ("Execution", "OrderIntent、Order、Fill、VenueState", "订单状态机与交易所同步"),
        ("Ledger", "CashEntry、PositionLot、Fee、FundingPayment、EquitySnapshot", "双重校验的资金事实"),
        ("Governance", "Approval、LimitVersion、ModelCard、AuditEvent", "审批、版本和审计"),
    ]
    add_table(doc, ["限界上下文", "聚合/实体", "职责"], domains, [1800, 3800, 3760], 9)

    add_heading(doc, "5 数据架构", 1)
    add_heading(doc, "5.1 分层", 2)
    add_table(doc, ["层级", "示例", "不可变性与版本"], [
        ("Raw", "交易所响应、WebSocket事件、新闻原文", "追加写；保存source_ts/receive_ts"),
        ("Normalized", "统一Tick、Bar、Funding、OI", "schema_version与质量标记"),
        ("Feature", "动量、波动、订单流、宏观状态", "feature_set_version与as_of_time"),
        ("Prediction", "概率、分位数、模型不确定度", "model/calibration/version"),
        ("Execution/Ledger", "订单、成交、费用、仓位、权益", "不可覆盖；用冲正记录修正"),
    ], [1800, 3800, 3760], 9.2)
    add_heading(doc, "5.2 数据质量", 2)
    add_bullets(doc, ["Completeness：周期内缺口率和连续缺口。", "Timeliness：source_ts 到 receive_ts 延迟及最后更新时间。", "Validity：价格、数量、时间、精度和枚举合法性。", "Consistency：跨源价格偏差、OHLC约束和成交量单调规则。", "Uniqueness：交易所事件ID或复合唯一键去重。"])
    add_heading(doc, "5.3 核心数据库调整", 2)
    add_table(doc, ["对象", "关键字段/约束"], [
        ("instrument", "venue、symbol、type、tick_size、lot_size、status、valid_time"),
        ("market_event", "source_ts、receive_ts、sequence、quality_status、raw_ref"),
        ("feature_value", "feature_set_id、as_of_time、value、source_window"),
        ("forecast", "horizon、p_up/p_down/p_flat、quantiles、model_version"),
        ("order", "client_order_id UNIQUE、venue_order_id、state、version"),
        ("ledger_entry", "account、currency、amount、entry_type、reference_id、immutable"),
        ("audit_event", "actor、action、resource、before/after、correlation_id"),
    ], [2300, 7060], 9.2)

    add_heading(doc, "6 量化预测设计", 1)
    add_p(doc, "预测目标采用未来对数收益 r(t,h)=ln(P(t+h)/P(t))。模型同时输出方向概率、条件收益分位数和波动率，不直接将上涨概率解释为预期收益。")
    add_table(doc, ["组件", "设计"], [
        ("Label", "按交易周期定义h和中性阈值θ；扣除可交易成本后标记方向。"),
        ("Feature", "价格、量价、订单流、衍生品、波动、跨资产、新闻/宏观/链上。"),
        ("Regime", "HMM/Change Point/聚类识别趋势、震荡、高波动和流动性危机。"),
        ("Baseline", "随机、永远看涨、动量、均值回复、逻辑回归。"),
        ("Model", "优先LightGBM/XGBoost；时序深度模型仅在基线后验证。"),
        ("Calibration", "Platt/Isotonic；按周期、品种和regime监控Brier/ECE。"),
        ("Validation", "Walk-forward、purged K-fold、embargo、成本压力测试。"),
    ], [2100, 7260], 9.2)
    add_note(doc, "决策函数", "EV = Σ p(state) × conditional_return(state) - fee - slippage - funding - impact。仅 EV、概率置信区间和 CVaR 同时达标时进入组合层。")

    add_heading(doc, "7 组合与风险设计", 1)
    add_table(doc, ["层级", "规则"], [
        ("信号", "最低数据质量、最低净期望、预测新鲜度、模型适用状态。"),
        ("仓位", "risk_budget / stop_distance，并以波动率和流动性缩放。"),
        ("组合", "单资产、相关簇、方向、策略、交易所、币种和总杠杆上限。"),
        ("尾部", "历史/参数CVaR、压力场景、跳空和流动性折价。"),
        ("行为", "连续亏损、频率、异常加仓、模型漂移触发降级。"),
        ("熔断", "阻止新订单、取消挂单、可选减仓；恢复必须审批。"),
    ], [1800, 7560], 9.3)

    add_heading(doc, "8 回测与执行设计", 1)
    add_heading(doc, "8.1 统一订单状态机", 2)
    add_p(doc, "CREATED → RISK_APPROVED → SUBMITTED → ACKNOWLEDGED → PARTIALLY_FILLED → FILLED；异常分支为 REJECTED、CANCEL_PENDING、CANCELLED、EXPIRED、UNKNOWN。UNKNOWN 必须通过交易所回查，不得盲目重试。")
    add_heading(doc, "8.2 撮合原则", 2)
    add_bullets(doc, ["市价单使用下一可交易事件价格并加入方向性滑点。", "限价触及不等于必然成交，应支持排队/成交概率模型。", "反手拆为平旧仓和开新仓两笔会计动作。", "手续费、滑点、资金费率和借贷成本分别入账。", "任何事件后执行现金、仓位、权益和PnL恒等式校验。"])
    add_heading(doc, "8.3 绩效", 2)
    add_bullets(doc, ["以日末权益计算收益，不使用日内峰值代理。", "分别报告毛收益、交易成本、资金成本与净收益。", "输出Sharpe、Sortino、Calmar、最大回撤、回撤时长、CVaR和换手率。", "按策略、Agent、品种、regime和时间归因。"])

    add_heading(doc, "9 Agent 与 LLM 设计", 1)
    add_table(doc, ["职责", "允许", "禁止"], [
        ("信息抽取", "新闻事件、实体、影响方向、时间范围", "执行新闻中的指令"),
        ("观点生成", "基于结构化证据解释支持/反对理由", "虚构缺失数据"),
        ("辩论", "暴露冲突、补充反证、标记不可用", "用角色数量制造假共识"),
        ("决策说明", "解释确定性评分器结果", "覆盖风控和订单方向"),
        ("自演进", "提出候选Prompt/策略版本并离线评估", "在线直接替换生产策略"),
    ], [1900, 3700, 3760], 9.2)
    add_p(doc, "所有 LLM 输出先按 JSON Schema 验证，再进入评分器。来源为 hardcoded、fallback 或 unavailable 的观点，其有效权重必须为零。Prompt 版本必须经过离线回放、样本外评估和人工审批。")

    add_heading(doc, "10 接口与事件设计", 1)
    add_table(doc, ["规范", "要求"], [
        ("REST", "统一错误码、request_id、分页、幂等键和版本前缀。"),
        ("WebSocket", "认证后订阅；频道级授权；sequence和重连补偿。"),
        ("内部事件", "event_id、type、version、occurred_at、correlation_id、payload。"),
        ("任务", "持久化状态、租约、attempt、next_retry_at、dead_letter。"),
        ("兼容", "新增字段向后兼容；破坏变更发布新API版本。"),
    ], [1900, 7460], 9.4)

    add_heading(doc, "11 前端设计", 1)
    add_bullets(doc, ["页面容器负责路由与用例，展示组件不直接访问API。", "统一 Server State 查询缓存，Pinia仅保存认证、偏好和跨页工作流。", "决策卡展示周期、概率、期望、成本、风险、证据、反证和失效条件。", "实盘、模拟、纸面环境使用持续可见且不可混淆的颜色和文字标识。", "高危按钮采用明确结果描述、二次确认和不可逆性说明。", "图表支持数据时间、来源、缺口和过期状态。"])

    add_heading(doc, "12 安全设计", 1)
    add_table(doc, ["威胁", "控制"], [
        ("越权访问", "所有权过滤+服务层授权+RLS+自动化越权测试"),
        ("密钥泄漏", "信封加密、轮换、脱敏、最小交易所权限、禁止提币"),
        ("订单重放", "幂等键、nonce/时间窗、订单状态回查"),
        ("Prompt注入", "外部文本隔离、结构化抽取、工具白名单、输出Schema"),
        ("供应链", "锁定依赖、SCA、镜像签名、SBOM"),
        ("审计篡改", "追加写审计、哈希链/外部归档、时间同步"),
    ], [2000, 7360], 9.3)

    add_heading(doc, "13 可观测性与运维", 1)
    add_bullets(doc, ["技术指标：请求延迟、错误、DB池、任务积压、WebSocket重连。", "数据指标：延迟、缺口、异常、跨源偏差、最后有效时间。", "模型指标：覆盖率、概率校准、漂移、各regime表现。", "交易指标：订单拒绝、未知状态、成交延迟、滑点、手续费、仓位差异。", "风险指标：暴露、回撤、CVaR、限额接近度和熔断状态。"])

    add_heading(doc, "14 测试策略", 1)
    add_table(doc, ["测试层", "重点"], [
        ("单元测试", "指标、仓位、费用、PnL、状态机、权限规则"),
        ("属性测试", "随机订单下资金守恒、仓位不为非法负值、概率和为1"),
        ("集成测试", "数据库事务、OKX适配、任务恢复、用户隔离"),
        ("回放测试", "固定原始事件产生可重复预测与订单"),
        ("压力测试", "行情突发、LLM超时、数据库降级、交易所未知订单"),
        ("前端E2E", "登录、分析、模拟、风险确认、紧停、越权与错误状态"),
    ], [1900, 7460], 9.3)

    add_heading(doc, "15 迁移与发布方案", 1)
    add_steps(doc, ["冻结当前回测晋级结果，建立问题基线。", "修复账本并对历史回测重新计算，旧结果标记为不可比较。", "引入统一领域接口，通过适配器兼容现有路由。", "建立影子预测，只记录不执行，完成概率校准。", "纸面盘→交易所模拟盘→小额实盘逐级发布。", "每级设置回滚、降级、熔断和人工审批门槛。"])
    path = OUT / "MoneyRobert-Pro_开发设计文档_V2.0.docx"
    doc.save(path)
    return path


def manual_document():
    doc = Document()
    style_document(doc, "MoneyRobert Pro · 系统操作手册")
    cover(doc, "系统操作手册", "市场分析、Agent 决策、模拟交易、回测、风控与系统管理", "MRP-UM-002")
    add_doc_control(doc, "指导用户安全完成系统配置、市场分析、预测阅读、模拟交易、回测、自动化授权、风险控制和故障处理。", "普通用户、交易员、基金经理、风控人员、系统管理员和运维人员")
    add_contents(doc, ["使用前须知", "登录与基础设置", "市场与行情", "AI分析与预测", "Agent辩论", "策略与回测", "纸面/模拟交易", "实盘交易", "自动交易与晋级", "风险与紧急停止", "报告与通知", "管理员操作", "常见问题与故障处理", "安全检查清单"])

    add_heading(doc, "1 使用前须知", 1)
    add_note(doc, "风险声明", "系统分析不构成收益承诺。当前版本在回测账本和绩效计算完成 P0 整改前，不应依据现有回测结果扩大实盘权限。", "risk")
    add_heading(doc, "1.1 环境识别", 2)
    add_table(doc, ["环境", "资金", "用途", "建议"], [
        ("纸面交易", "系统虚拟资金", "流程与策略初验", "默认使用"),
        ("OKX模拟盘", "交易所模拟资金", "验证订单和交易所适配", "纸面盘通过后使用"),
        ("实盘", "真实资金", "受控执行", "仅在全部上线门槛通过后启用"),
    ], [1800, 1900, 3300, 2360], 9.2)
    add_bullets(doc, ["操作前确认页面顶部环境标识、账户和交易对。", "不要使用拥有提币权限的交易所 API Key。", "先设置最大可接受损失，再讨论收益目标。", "置信度不是保证概率；同时查看样本量、历史校准和数据时间。"])

    add_heading(doc, "2 登录与基础设置", 1)
    add_heading(doc, "2.1 注册与登录", 2)
    add_steps(doc, ["进入注册页，填写用户名、邮箱和高强度密码。", "登录后检查用户名和角色是否正确。", "长时间无人操作或设备丢失时退出全部会话。"])
    add_heading(doc, "2.2 配置交易所 API Key", 2)
    add_steps(doc, ["在 OKX 创建独立 API Key，仅勾选读取和交易权限，禁止提币。", "打开“系统设置 → API Key”，选择 OKX 与模拟盘/实盘类型。", "填写 Key、Secret、Passphrase 并保存。", "执行连接测试，确认账户、权限和环境一致。"])
    add_note(doc, "实盘保护", "首次配置必须先使用模拟盘密钥。若连接结果、账户或环境无法确认，请停止后续操作。", "warn")
    add_heading(doc, "2.3 配置 AI 服务", 2)
    add_steps(doc, ["进入“系统设置 → AI服务商”。", "填写供应商、模型、Base URL 和 API Key。", "将温度设为低到中等，保证结构化输出稳定。", "执行测试调用并检查用量限制。"])

    add_heading(doc, "3 市场与行情", 1)
    add_steps(doc, ["打开“市场行情”，选择交易对和K线周期。", "检查最新价、买卖价、成交量和数据更新时间。", "查看资金费率、持仓量和多空比；确认数据来源是实时 OKX 还是数据库缓存。", "若显示过期、缺口或无数据，不进入交易流程。"])
    add_table(doc, ["指标", "含义", "常见误区"], [
        ("资金费率", "永续多空资金交换与拥挤度线索", "正费率不等于马上下跌"),
        ("持仓量", "未平仓合约规模", "单看增加无法判断方向"),
        ("多空比", "账户或仓位的相对比例", "不同口径不能直接比较"),
        ("RSI/MACD", "动量与趋势特征", "超买超卖不等于反转"),
        ("ATR", "波动幅度", "不是方向预测"),
    ], [1700, 3500, 4160], 9.2)

    add_heading(doc, "4 AI 分析与预测", 1)
    add_steps(doc, ["打开“AI分析”，选择交易对和预测周期。", "分别查看技术、资金、情绪和综合分析。", "确认输入数据更新时间和缺失项。", "阅读建议动作、风险等级、止损/止盈和理由。", "在模拟盘验证，不从分析页直接扩大实盘仓位。"])
    add_heading(doc, "4.1 正确阅读置信度", 2)
    add_bullets(doc, ["当前 Agent 置信度表示系统意见强度，不天然等于真实上涨概率。", "优先查看历史相似样本量、校准表现和不同市场状态下的表现。", "Agent 高度分歧、数据缺失或市场状态异常时，应选择观望。", "任何建议都应扣除手续费、滑点和资金费率后再判断。"])

    add_heading(doc, "5 Agent 辩论", 1)
    add_steps(doc, ["进入“AI辩论分析”，选择交易对。", "启动辩论，等待技术、资金和新闻部门完成独立意见。", "查看各 Agent 的来源、方向、置信度与分析内容。", "重点阅读部门冲突和反对证据。", "查看基金经理最终动作、交易计划和风控否决原因。", "将不可用、硬编码兜底或数据缺失的 Agent 视为无效证据。"])
    add_note(doc, "判断原则", "Agent 数量多不代表证据更强。若多个 Agent 使用同一数据或同类指标，它们不是独立样本。")

    add_heading(doc, "6 策略与回测", 1)
    add_heading(doc, "6.1 创建策略", 2)
    add_steps(doc, ["进入“策略管理”，新建策略。", "选择交易对、方向、策略类型和参数。", "设置止损、止盈、最大仓位、最大杠杆和有效期。", "保存后先保持停用状态，进入回测。"])
    add_heading(doc, "6.2 运行回测", 2)
    add_steps(doc, ["选择不参与调参的历史区间。", "填写初始资金、手续费、滑点和资金费率假设。", "启动回测并记录数据版本与策略版本。", "查看净收益、最大回撤、样本量、盈亏比、成本和不同市场状态表现。", "使用样本外区间复验，不按单次最高收益选择参数。"])
    add_table(doc, ["必须查看", "最低解释要求"], [
        ("交易数", "样本是否足够，是否集中在少数日期"),
        ("最大回撤", "回撤幅度、持续时间和恢复时间"),
        ("净收益", "已扣手续费、滑点、资金费率"),
        ("Sharpe/Sortino", "计算频率、无风险利率和收益序列口径"),
        ("分状态表现", "趋势、震荡、高波动环境是否一致"),
    ], [2500, 6860], 9.3)
    add_note(doc, "当前限制", "现有撮合账本完成 P0 修复前，回测仅用于流程验证，不作为实盘晋级证据。", "risk")

    add_heading(doc, "7 纸面与模拟交易", 1)
    add_steps(doc, ["进入“纸面交易”或 Agent 仪表盘。", "选择交易对、初始虚拟资金和分析周期。", "启动模拟并观察首次决策，确认没有立即异常下单。", "检查持仓、未实现盈亏、已实现盈亏、费用和权益。", "对照交易历史确认每次开仓、加减仓和平仓原因。", "达到预设观察期后生成报告，不只看胜率。"])
    add_heading(doc, "7.1 模拟通过门槛", 2)
    add_bullets(doc, ["足够交易样本和足够长的运行时间。", "最大回撤、单日亏损和连续亏损均未突破限制。", "净收益为正且不是由单笔交易贡献。", "系统重启、断网和行情异常演练通过。", "账本与交易所模拟账户核对一致。"])

    add_heading(doc, "8 实盘交易", 1)
    add_note(doc, "启用条件", "只有风险负责人确认 P0 整改、模拟验证、权限隔离、订单幂等和紧急停止演练全部通过后，才可使用实盘。", "risk")
    add_steps(doc, ["确认当前环境明确显示“实盘”和正确账户。", "检查可用余额、现有仓位和总风险暴露。", "输入订单方向、类型、数量；限价单填写价格。", "设置止损和止盈，查看预估最大损失和全部成本。", "复核交易对、方向、杠杆、数量和环境后确认。", "下单后等待交易所确认；UNKNOWN 状态不得重复下单。"])

    add_heading(doc, "9 自动交易与晋级", 1)
    add_table(doc, ["等级", "模式", "要求"], [
        ("L0", "内部纸面", "验证基本决策和账本"),
        ("L1", "交易所模拟", "验证交易所订单与同步"),
        ("L2", "受控实盘", "小额、低杠杆、人工确认"),
        ("L3", "受限自主", "长期样本、双重审批、观察期和熔断"),
    ], [1300, 2600, 5460], 9.4)
    add_steps(doc, ["在 Agent 仪表盘查看当前等级和未满足条件。", "完成模拟交易、风险指标和运行天数要求。", "提交晋级申请，等待系统预检和人工审核。", "签署风险确认并设置最大可接受亏损。", "观察期内使用减半仓位和更低频率。", "出现模型漂移、越限或异常时自动降级。"])

    add_heading(doc, "10 风险与紧急停止", 1)
    add_heading(doc, "10.1 日常检查", 2)
    add_bullets(doc, ["单笔风险、总仓位、总杠杆和相关资产集中度。", "日/周回撤、连续亏损和交易频率。", "数据新鲜度、模型适用状态和预测覆盖率。", "交易所连接、订单未知状态和账户差异。"])
    add_heading(doc, "10.2 紧急停止", 2)
    add_steps(doc, ["点击 Agent 仪表盘右上角“紧急停止”。", "确认停止新订单，并选择是否取消挂单。", "检查现有仓位，按预案决定保持、减仓或全部平仓。", "记录触发原因和账户快照。", "问题排除后由授权人员审批恢复，不直接重新启动。"])

    add_heading(doc, "11 报告与通知", 1)
    add_bullets(doc, ["报告至少包含期间、策略版本、模型版本、净收益、回撤、成本、样本量和异常事件。", "通知按信息、警告、高风险和紧急四级分类。", "紧急通知必须包含账户、交易对、风险类型、当前处置和下一步。", "报告只允许所有者或授权角色访问、修改和导出。"])

    add_heading(doc, "12 管理员操作", 1)
    add_table(doc, ["任务", "操作要点"], [
        ("用户与角色", "最小权限；高权限定期复核；禁用离职/异常账户"),
        ("数据源", "配置采集频率、保留期；检查最后有效时间和缺口"),
        ("AI供应商", "控制模型、配额、超时和密钥；记录调用审计"),
        ("风险参数", "版本化、双人审批；禁止无记录放宽"),
        ("模型发布", "模型卡、样本外报告、影子期、回滚版本"),
        ("系统健康", "数据库、Redis、任务、行情、交易所和磁盘容量"),
    ], [2300, 7060], 9.3)

    add_heading(doc, "13 常见问题与故障处理", 1)
    add_table(doc, ["现象", "检查与处理"], [
        ("没有行情", "检查网络/代理、OKX状态、采集任务、最后数据时间；禁止使用旧缓存交易。"),
        ("AI分析失败", "检查供应商密钥、配额、模型名、代理和响应Schema；不得用伪造意见替代。"),
        ("订单状态未知", "停止重复提交，使用订单ID回查交易所；必要时人工对账。"),
        ("账户权益不一致", "立即停止自动交易，核对成交、费用、资金费率和账本事件。"),
        ("WebSocket断开", "自动重连后检查sequence缺口并REST补偿。"),
        ("回测异常高收益", "检查未来函数、费用、反手、现金记账、样本区间和幸存者偏差。"),
        ("频繁触发熔断", "保持停用，检查策略失效、波动状态、数据异常和风险参数。"),
    ], [2500, 6860], 9.2)

    add_heading(doc, "14 安全检查清单", 1)
    add_bullets(doc, ["API Key 无提币权限，模拟和实盘密钥严格分离。", "页面环境、账户、交易对、方向、数量和杠杆均已复核。", "数据更新时间和来源正常，无缺口或降级。", "最大损失、止损、组合暴露和熔断阈值已经设置。", "报告和资源只能被授权用户访问。", "紧急停止已演练，联系人和恢复审批人明确。", "系统、依赖、数据库和备份状态正常。"])
    path = OUT / "MoneyRobert-Pro_系统操作手册_V2.0.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for generated in (requirement_document(), design_document(), manual_document()):
        print(generated)
