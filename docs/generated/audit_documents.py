from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


root = Path(__file__).resolve().parent / "final"
failures = []

for path in sorted(root.glob("*.docx")):
    with ZipFile(path) as zf:
        bad = zf.testzip()
        if bad:
            failures.append(f"{path.name}: corrupt member {bad}")

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    if len(text.strip()) < 6000:
        failures.append(f"{path.name}: content unexpectedly short")
    if not headings:
        failures.append(f"{path.name}: no headings")
    if any(not p.text.strip() for p in headings):
        failures.append(f"{path.name}: empty heading")
    for token in ("TODO", "TBD", "placeholder", "Lorem ipsum"):
        if token in text:
            failures.append(f"{path.name}: placeholder token {token}")

    for sec in doc.sections:
        if abs(sec.page_width - 7772400) > 100 or abs(sec.page_height - 10058400) > 100:
            failures.append(f"{path.name}: page geometry is not Letter")
        for margin in (sec.top_margin, sec.right_margin, sec.bottom_margin, sec.left_margin):
            if abs(margin - 914400) > 100:
                failures.append(f"{path.name}: margin is not 1 inch")

    for ti, table in enumerate(doc.tables, 1):
        grid = table._tbl.tblGrid
        widths = [int(c.get(qn("w:w"))) for c in grid]
        if not widths or sum(widths) != 9360:
            failures.append(f"{path.name}: table {ti} grid width {sum(widths) if widths else 0}")
        for ri, row in enumerate(table.rows, 1):
            if len(row.cells) != len(widths):
                failures.append(f"{path.name}: table {ti} row {ri} column mismatch")
            for ci, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tcw is None or int(tcw.get(qn("w:w"))) != widths[ci]:
                    failures.append(f"{path.name}: table {ti} row {ri} cell width mismatch")
                trh = row._tr.get_or_add_trPr().find(qn("w:trHeight"))
                if trh is not None and trh.get(qn("w:hRule")) == "exact":
                    failures.append(f"{path.name}: table {ti} has fixed row height")

    print(f"{path.name}: chars={len(text)} headings={len(headings)} tables={len(doc.tables)} package=OK")

if failures:
    print("FAIL")
    for f in failures:
        print(f)
    raise SystemExit(1)
print("STRUCTURAL AUDIT PASSED")
